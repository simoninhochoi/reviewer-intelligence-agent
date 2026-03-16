"""DOCX 파싱 유틸리티 (python-docx 기반) — 원고 텍스트 추출"""
from pathlib import Path

from docx import Document

from src.utils.logging_config import logger


def extract_text_from_docx(docx_path: Path | str) -> str:
    """DOCX 파일에서 전체 텍스트 추출

    Args:
        docx_path: DOCX 파일 경로

    Returns:
        추출된 전체 텍스트
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        logger.error(f"DOCX 파일을 찾을 수 없습니다: {docx_path}")
        return ""

    try:
        doc = Document(str(docx_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        logger.info(f"DOCX 텍스트 추출 완료: {docx_path.name} ({len(full_text)} chars)")
        return full_text
    except Exception as e:
        logger.error(f"DOCX 텍스트 추출 실패 ({docx_path}): {e}")
        return ""


def extract_sections_from_docx(docx_path: Path | str) -> dict[str, str]:
    """DOCX에서 헤딩 기반 섹션 분리

    Returns:
        {"Introduction": "...", "Methods": "...", ...} 형태의 섹션 딕셔너리
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        return {}

    try:
        doc = Document(str(docx_path))
        sections: dict[str, str] = {}
        current_heading = "Untitled"
        current_content: list[str] = []

        for para in doc.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                # 이전 섹션 저장
                if current_content:
                    sections[current_heading] = "\n\n".join(current_content)
                current_heading = para.text.strip() or "Untitled"
                current_content = []
            elif para.text.strip():
                current_content.append(para.text)

        # 마지막 섹션 저장
        if current_content:
            sections[current_heading] = "\n\n".join(current_content)

        return sections
    except Exception as e:
        logger.error(f"DOCX 섹션 추출 실패 ({docx_path}): {e}")
        return {}
